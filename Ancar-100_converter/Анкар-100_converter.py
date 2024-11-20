import tkinter as tk
import warnings
import os
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK_TYPE

class AnkarConverter(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Анкар-100 конвертер")
        self.geometry("400x400")
        self.configure(bg="#171a21")

        style = ttk.Style()
        style.theme_use("default")
        style.configure("TEntry", padding=8, font=("Arial", 12))
        style.configure("TButton", padding=8, font=("Arial", 12), background="#1b2838", foreground="white")
        style.map("TButton",
                  background=[("active", "#006400")],
                  foreground=[("active", "white")])

        self.file_path = tk.StringVar()
        self.file_entry = ttk.Entry(self, textvariable=self.file_path)
        self.file_entry.pack(pady=10)

        self.browse_button = ttk.Button(self, text="Выбрать файл", command=self.choose_file)
        self.browse_button.pack(pady=10)

        self.convert_button = ttk.Button(self, text="Перевести", command=self.convert_file)
        self.convert_button.pack(pady=10)

    def choose_file(self):
        file_path = filedialog.askopenfilename()
        self.file_path.set(file_path)
        if file_path:
            self.file_name = os.path.splitext(os.path.basename(file_path))[0]
        else:
            self.file_name = ""

    def convert_file(self):
        file_path = self.file_path.get()
        if file_path:
            try:
                if file_path.endswith(".xlsx"):
                    with warnings.catch_warnings():
                        warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")
                        self.convert_to_docx(file_path, self.file_name)
                        tk.messagebox.showinfo("Успех", "Файл конвертирован")
                else:
                    tk.messagebox.showerror("Ошибка", "Неподдерживаемый тип файла. Пожалуйста, выберите файл .xlsx.")
            except Exception as e:
                tk.messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")
                print(e)

    def convert_to_docx(self, file_path, file_name):
        # Create a new Word document
        document = Document()

        # Add headings
        document.add_heading('ОАО "Лидский молочно-консервный комбинат"', 0)
        document.add_heading('Журнал по входному контролю молока на наличие антибиотиков', 1)
        document.add_heading(' ', 2)

        # Read data from the XLSX file
        try:
            workbook = load_workbook(filename=file_path)
            sheet = workbook.active
            rows = list(sheet.iter_rows(values_only=True))
        except FileNotFoundError:
            print(f"Ошибка: Файл {file_path} не найден.")
            return
        except Exception as e:
            print(f"Ошибка при чтении XLSX файла: {e}")
            return

        # Add the table to the document
        try:
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
        except IndexError:
            print("Ошибка: Пустой XLSX файл или неверный формат данных.")
            return

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'TableGrid'

        for i, row in enumerate(rows):
            for j, cell in enumerate(table.rows[i].cells):
                cell.text = str(row[j])

        section = document.sections[0]
        section.page_width = Inches(28) 
        section.page_height = Inches(30)
        section.top_margin = Inches(1)

        # Save the document
        document.save(f'{file_name}.docx')


if __name__ == "__main__":
    app = AnkarConverter()
    app.mainloop()